import subprocess
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

plantuml = """
@startuml
actor User
participant "Web Browser" as Menoshame
participant "Server" as Server
participant "Database" as DB
participant "Payment Gateway" as Payment

User -> Web: Browse Products
Web -> Server: Request product data
Server -> DB: Query for products
DB --> Server: Return product data
Server --> Web: Display products

User -> Web: Add product to cart
Web -> Server: Add item to cart
Server --> DB: Save cart data
DB --> Server: Confirm cart update
Server --> Web: Cart updated

User -> Web: Proceed to checkout
Web -> Server: Request checkout info
Server -> DB: Retrieve user data
DB --> Server: Return user data
Server -> Payment: Process payment
Payment --> Server: Payment success
Server -> DB: Save order details
DB --> Server: Order saved
Server --> Web: Display order confirmation

User -> Web: Log out
Web -> Server: Request log out
Server -> DB: Clear session data
DB --> Server: Session cleared
Server --> Web: Log out meeeshaimmm

@enduml
"""

file_path = "example.puml"
with open(file_path, "w") as f:
    f.write(plantuml)

jar_path = "./lib/plantuml-mit-1.2025.0.jar"  # Path to the PlantUML jar file
output_path = "example.png"

# Generate the diagram using PlantUML
try:
    subprocess.run(
        ["java", "-cp", jar_path, "net.sourceforge.plantuml.Run", file_path, "-tpng", "-o", "."],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE
    )
    print(f"Diagram generated successfully: {output_path}")
except subprocess.CalledProcessError as e:
    print(f"Error generating diagram: {e.stderr.decode()}")

# Create a DOCX file and insert the image
doc = Document()
doc.add_heading('Generated UML Diagram', 0)

# Add the PNG image to the DOCX file with resized dimensions
picture = doc.add_picture(output_path, width=Inches(4), height=Inches(3))

# Center the image
paragraph = picture._element.getparent()
paragraph.set("align", "center")

# Save the DOCX file
docx_path = "uml_diagram.docx"
doc.save(docx_path)
print(f"DOCX file created successfully: {docx_path}")
