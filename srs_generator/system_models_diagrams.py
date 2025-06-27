import subprocess
import os
from .rag import AgentBase
from docx.shared import Inches
from loguru import logger
import re
from dotenv import load_dotenv

class SystemModelsAgent(AgentBase):
    def __init__(self, max_retries=5, verbose=True):  # Increased retries for robustness
        super().__init__(name="SystemModelsAgent", max_retries=max_retries, verbose=verbose)
        load_dotenv()
        self.logger = logger
        self.diagram_types = {
            "ActivityDiagram": """Generate a detailed PlantUML Activity Diagram that shows the complete flow of actions for these use cases and with no errors.
                                Context: You are an expert UML designer. Create a comprehensive activity diagram that captures all workflows.
                                
                                Requirements:
                                1. Start with initialization nodes
                                2. Include all decision points and branches
                                3. Show parallel processes using fork/join when needed
                                4. End with proper termination nodes
                                5. Use swimlanes if multiple actors are involved
                                
                                Format Rules:
                                1. Use correct PlantUML syntax
                                2. Include start and stop nodes
                                3. Proper indentation for readability
                                4. Clear labels for all actions
                                5. Consistent arrow usage
                                
                                Example Syntax:
                                @startuml
                                |User|
                                start
                                :Login to System;
                                if (Valid Credentials?) then (yes)
                                    :Access Dashboard;
                                else (no)
                                    :Show Error;
                                    stop
                                endif
                                |System|
                                :Process Request;
                                stop
                                @enduml""",
            
            "SequenceDiagram": """Generate a detailed PlantUML Sequence Diagram showing interactions between components and without any error.
                                Context: You are an expert in system design. Create a sequence diagram showing all interactions without error.
                                
                                Requirements:
                                1. Include all system components and actors
                                2. Show complete message flow
                                3. Include alternative flows
                                4. Add activation bars
                                5. Show return messages
                                
                                Format Rules:
                                1. Use proper PlantUML syntax
                                2. Clear participant definitions
                                3. Consistent arrow types
                                4. Proper message labels
                                5. Include activation/deactivation
                                
                                Example Syntax:
                                @startuml
                                actor User
                                participant "Frontend" as FE
                                participant "Backend" as BE
                                database DB
                                
                                User -> FE: Login Request
                                activate FE
                                FE -> BE: Validate
                                activate BE
                                BE -> DB: Query
                                DB --> BE: Result
                                BE --> FE: Response
                                deactivate BE
                                FE --> User: Show Dashboard
                                deactivate FE
                                @enduml""",

            "ClassDiagram": """Generate a comprehensive PlantUML Class Diagram showing system structure and without any error.
                             Context: You are an expert in object-oriented design. Create a detailed class diagram without any error.
                             
                             Requirements:
                             1. Include all major classes
                             2. Show inheritance relationships
                             3. Include aggregation/composition
                             4. Show multiplicity
                             5. Include key methods and attributes
                             
                             Format Rules:
                             1. Use proper PlantUML syntax
                             2. Clear visibility modifiers
                             3. Proper relationship symbols
                             4. Type definitions for attributes
                             5. Method signatures with parameters
                             
                             Example Syntax:
                             @startuml
                             class User {
                                 -id: Long
                                 -email: String
                                 +login(credentials: Auth): boolean
                                 #validateInput(): void
                             }
                             
                             class Order {
                                 -items: List<Item>
                                 +calculateTotal(): decimal
                             }
                             
                             User "1" *-- "*" Order: places
                             @enduml"""
        }
        self.diagrams = {}

    def validate_diagram_code(self, code, diagram_type):
        """Validate the generated PlantUML code using Gemini"""
        if not code:
            self.logger.error(f"[{self.name}] No PlantUML code provided for {diagram_type}")
            return False

        validation_prompt = f"""You are a PlantUML expert validator. Review this {diagram_type} code for correctness.
        
        Code to validate:
        {code}
        
        Check for:
        1. Syntax correctness
        2. Proper start/end tags
        3. Logical flow and connections
        4. Required elements for {diagram_type}
        5. Proper relationship definitions
        
        Respond with ONLY 'VALID' or 'INVALID' followed by a brief reason."""
        
        messages = [
            self.format_message("system", validation_prompt),
            self.format_message("user", "Please validate the provided PlantUML code.")
        ]

        for attempt in range(self.max_retries):
            validation_result = self.call_gemini(messages, temperature=0.3, max_tokens=500)
            if validation_result and 'VALID' in validation_result.upper():
                self.logger.info(f"[{self.name}] {diagram_type} code validated successfully")
                return True
            self.logger.warning(f"[{self.name}] Validation attempt {attempt + 1} failed: {validation_result or 'No response'}")
        
        self.logger.error(f"[{self.name}] Validation failed for {diagram_type} after {self.max_retries} attempts")
        return False

    def extract_plantuml(self, text):
        """Extract PlantUML code between @startuml and @enduml tags"""
        if not text:
            self.logger.error(f"[{self.name}] No text provided to extract PlantUML code")
            return None
        
        match = re.search(r'(@startuml[\s\S]*?@enduml)', text, re.DOTALL)
        if not match:
            self.logger.error(f"[{self.name}] No valid PlantUML code found in the response")
            return None
        
        code = match.group(1).strip()
        # Preserve newlines and formatting while removing extra whitespace
        code = '\n'.join(line.strip() for line in code.split('\n'))
        return code

    def generate_diagram_code(self, diagram_type, topic, previous_contents):
        """Generate PlantUML code using Gemini with improved prompting"""
        if 'use_cases' not in previous_contents:
            self.logger.error(f"[{self.name}] No use cases found in previous contents")
            return None

        system_message = self.diagram_types.get(diagram_type, "")
        user_message = (
            f"Here is the project description:\n{topic}\n\n"
            f"Here is the introduction:\n{previous_contents['introduction']}\n\n"
            f"Here is the overall description:\n{previous_contents['overall_description']}\n\n"
            f"Here are the system features:\n{previous_contents['system_features']}\n\n"
            f"Here are the external interfaces:\n{previous_contents['external_interfaces']}\n\n"
            f"Here are the non-functional requirements:\n{previous_contents['non_functional_requirements']}\n\n"
            f"Here are the use cases:\n{previous_contents['use_cases']}\n\n"
            f"Generate ONLY the PlantUML code for a {diagram_type}, ensuring all elements are properly connected, with clear labels, and no explanations."
        )

        messages = [
            self.format_message("system", system_message),
            self.format_message("user", user_message)
        ]

        for attempt in range(self.max_retries):
            try:
                response = self.call_gemini(messages, temperature=0.3, max_tokens=2000)
                if not response:
                    self.logger.error(f"[{self.name}] No response from Gemini API for {diagram_type} in attempt {attempt + 1}")
                    continue
                
                plantuml_code = self.extract_plantuml(response)
                if not plantuml_code:
                    self.logger.error(f"[{self.name}] Failed to extract PlantUML code for {diagram_type} in attempt {attempt + 1}")
                    continue
                
                if self.validate_diagram_code(plantuml_code, diagram_type):
                    self.logger.info(f"[{self.name}] Successfully generated valid PlantUML code for {diagram_type}")
                    return plantuml_code
                
                self.logger.warning(f"[{self.name}] Attempt {attempt + 1}: Generated code failed validation for {diagram_type}")
            except Exception as e:
                self.logger.error(f"[{self.name}] Error in attempt {attempt + 1} for {diagram_type}: {str(e)}")
        
        self.logger.error(f"[{self.name}] Failed to generate valid PlantUML code for {diagram_type} after {self.max_retries} attempts")
        return None

    def create_diagram(self, name, plantuml_code):
        """Generate PNG from PlantUML code with robust error handling"""
        if not plantuml_code:
            self.logger.error(f"[{self.name}] No PlantUML code provided for {name}")
            return None
            
        try:
            # Create diagrams directory if it doesn't exist
            os.makedirs("diagrams", exist_ok=True)
            
            # Clean and normalize the PlantUML code
            plantuml_code = self._normalize_plantuml_code(plantuml_code)
            
            # Save PUML file with proper newlines and encoding
            puml_file = os.path.join("diagrams", f"{name}.puml")
            with open(puml_file, "w", encoding='utf-8', newline='\n') as f:
                f.write(plantuml_code)
                
            # Verify the saved file
            if not os.path.exists(puml_file):
                self.logger.error(f"[{self.name}] Failed to save PUML file: {puml_file}")
                return None
                
            # Find PlantUML jar with expanded search paths
            jar_name = "plantuml-mit-1.2025.0.jar"
            jar_paths = [
                os.path.join(".", "lib", jar_name),
                os.path.join(os.path.dirname(__file__), "lib", jar_name),
                os.path.join(os.path.dirname(os.path.dirname(__file__)), "lib", jar_name),
                os.path.abspath(os.path.join("lib", jar_name))
            ]
            
            jar_path = None
            for path in jar_paths:
                if os.path.exists(path):
                    jar_path = path
                    break
                    
            if not jar_path:
                self.logger.error(f"[{self.name}] PlantUML jar not found in paths: {jar_paths}")
                return None
                
            # Set up Java command with explicit options
            cmd = [
                "java",
                "-Djava.awt.headless=true",  # Run in headless mode
                "-DPLANTUML_LIMIT_SIZE=8192",  # Increase size limit
                "-jar",
                jar_path,
                "-charset",
                "UTF-8",
                "-tpng",
                "-timeout",
                "60",  # Set timeout in seconds
                puml_file
            ]
            
            # Log the exact file contents being processed
            self.logger.debug(f"[{self.name}] Processing PlantUML file: {puml_file}")
            self.logger.debug(f"[{self.name}] File contents:")
            with open(puml_file, 'r', encoding='utf-8') as f:
                self.logger.debug(f.read())
                
            # Run PlantUML with explicit error handling
            try:
                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    check=False  # Don't raise exception on non-zero exit
                )
                
                # Always log command output for debugging
                if result.stdout:
                    self.logger.debug(f"[{self.name}] PlantUML stdout: {result.stdout}")
                if result.stderr:
                    self.logger.debug(f"[{self.name}] PlantUML stderr: {result.stderr}")
                    
                # Check specific error codes
                if result.returncode == 200:
                    self.logger.error(f"[{self.name}] PlantUML syntax error detected")
                    # Try to fix common syntax issues
                    fixed_code = self._fix_common_plantuml_issues(plantuml_code)
                    if fixed_code != plantuml_code:
                        self.logger.info(f"[{self.name}] Attempting with fixed PlantUML code")
                        return self.create_diagram(name, fixed_code)
                elif result.returncode != 0:
                    self.logger.error(f"[{self.name}] PlantUML execution failed with code: {result.returncode}")
                    return None
                    
                # Verify PNG was generated
                png_path = os.path.join("diagrams", f"{name}.png")
                if os.path.exists(png_path):
                    self.logger.info(f"[{self.name}] Successfully generated {name} image at {png_path}")
                    return png_path
                else:
                    self.logger.error(f"[{self.name}] PNG file not generated despite successful execution")
                    return None
                    
            except Exception as e:
                self.logger.error(f"[{self.name}] Error executing PlantUML: {str(e)}")
                return None
                
        except Exception as e:
            self.logger.error(f"[{self.name}] Error in create_diagram: {str(e)}")
            return None
    
    def _normalize_plantuml_code(self, code):
        """Normalize PlantUML code to prevent common issues"""
        if not code:
            return None
            
        # Ensure proper line endings
        code = code.replace('\r\n', '\n').replace('\r', '\n')
        
        # Clean up the code
        lines = code.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Remove excess whitespace but preserve indentation
            line = line.rstrip()
            
            # Ensure proper arrow syntax
            line = line.replace('-->', '-->').replace('->', '->')
            
            cleaned_lines.append(line)
            
        code = '\n'.join(cleaned_lines)
        
        # Ensure proper start/end tags
        if not code.startswith('@startuml'):
            code = '@startuml\n' + code
        if not code.endswith('@enduml'):
            code = code + '\n@enduml'
            
        return code
        
    def _fix_common_plantuml_issues(self, code):
        """Fix common PlantUML syntax issues"""
        if not code:
            return None
            
        # Common fixes
        fixes = [
            (r'(?<![->])->', '-->'),  # Fix arrow syntax
            (r'\bend\b(?!if|while|fork)', 'end'),  # Fix end statements
            (r'(\w+)\s*-+\s*(\w+)', r'\1 --> \2'),  # Fix relationship arrows
            (r'(\w+)\s*=+\s*(\w+)', r'\1 == \2'),  # Fix double lines
            (r'(?<!@)start\b', '@startuml'),  # Fix missing @ in start
            (r'(?<!@)end\b(?!if|while|fork)', '@enduml')  # Fix missing @ in end
        ]
        
        fixed_code = code
        for pattern, replacement in fixes:
            fixed_code = re.sub(pattern, replacement, fixed_code)
            
        # Ensure proper start/end tags
        if not fixed_code.startswith('@startuml'):
            fixed_code = '@startuml\n' + fixed_code
        if not fixed_code.endswith('@enduml'):
            fixed_code = fixed_code + '\n@enduml'
            
        return fixed_code
    
    def add_diagrams_to_doc(self, doc, topic, previous_contents):
        """Generate diagrams and add them to Word document without duplication"""
        # Check if section already exists
        section_exists = False
        for paragraph in doc.paragraphs:
            if paragraph.text == "7. System Models and Diagrams":
                section_exists = True
                break
        
        if not section_exists:
            # Add new section if it doesn't exist
            doc.add_heading("7. System Models and Diagrams", level=1)
            doc.add_paragraph("This section presents the system models using various UML diagrams to visualize different aspects of the system.")
        else:
            # If section exists, add spacing
            doc.add_paragraph()

        # Generate and add each diagram
        for index, diagram_type in enumerate(self.diagram_types, 1):
            # Check if this diagram type already exists
            diagram_exists = False
            for paragraph in doc.paragraphs:
                if paragraph.text == f"7.{index} {diagram_type}":
                    diagram_exists = True
                    break
            
            if diagram_exists:
                self.logger.info(f"[{self.name}] {diagram_type} already exists in document, skipping...")
                continue
                
            # Generate PlantUML code using LLM
            plantuml_code = self.generate_diagram_code(diagram_type, topic, previous_contents)
            if not plantuml_code:
                self.logger.warning(f"[{self.name}] Failed to generate PlantUML code for {diagram_type}, skipping...")
                doc.add_heading(f"7.{index} {diagram_type}", level=2)
                doc.add_paragraph(f"Failed to generate {diagram_type} image.")
                continue
                
            # Store generated code
            self.diagrams[diagram_type] = plantuml_code
                
            # Create diagram
            png_path = self.create_diagram(diagram_type, plantuml_code)
            if not png_path:
                self.logger.warning(f"[{self.name}] Failed to create diagram image for {diagram_type}, adding placeholder...")
                doc.add_heading(f"7.{index} {diagram_type}", level=2)
                doc.add_paragraph(f"Failed to generate {diagram_type} image.")
                continue
                
            # Add to document
            doc.add_heading(f"7.{index} {diagram_type}", level=2)
            doc.add_picture(png_path, width=Inches(6))
            doc.add_paragraph()  # Add spacing

    def execute(self, topic, previous_contents, file_name):
        """Main execution method required by AgentBase"""
        if 'use_cases' not in previous_contents:
            self.logger.error(f"[{self.name}] No use cases found in previous contents")
            return previous_contents

        try:
            # Load existing document
            from docx import Document
            doc = Document(file_name) if os.path.exists(file_name) else Document()
            
            # Generate and add diagrams
            self.add_diagrams_to_doc(doc, topic, previous_contents)
            
            # Save document
            doc.save(file_name)
            
            # Add generated diagrams to previous contents
            previous_contents['system_models'] = self.diagrams
            
        except Exception as e:
            self.logger.error(f"[{self.name}] Error in execute: {str(e)}")
            
        return previous_contents