# first_page.py
import os
from abc import ABC, abstractmethod
from loguru import logger
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import time

class SRSBase(ABC):
    def __init__(self, name, max_retries=2, verbose=True):
        self.name = name
        self.max_retries = max_retries
        self.verbose = verbose
        self.doc = Document()
        self.headings = []  # Track headings for TOC

    @abstractmethod
    def create_first_page(self, user_name, file_name):
        pass

    @abstractmethod
    def add_page(self, content, file_name):
        pass

    def save_document(self, file_name):
        retries = 0
        while retries < self.max_retries:
            try:
                if self.verbose:
                    logger.info(f"[{self.name}] Saving document as {file_name}")
                self.doc.save(file_name)
                logger.info(f"[{self.name}] Document saved successfully.")
                return
            except Exception as e:
                retries += 1
                logger.error(f"[{self.name}] Error saving document: {e}. Retry {retries}/{self.max_retries}")
        raise Exception(f"[{self.name}] Failed to save document after {self.max_retries} retries.")

class SRSConcrete(SRSBase):
    def __init__(self, name, max_retries=2, verbose=True):
        super().__init__(name, max_retries, verbose)
        self.bold_pattern = re.compile(r'\*\*(.*?)\*\*|\*(.*?)\*')
        self.heading_pattern = re.compile(r'^#+\s+(.*?)$', re.MULTILINE)
        self.bullet_pattern = re.compile(r'^\s*[-*]\s+(.*?)$', re.MULTILINE)

    def add_table_of_contents(self):
        """Insert a Table of Contents field at the beginning of the document."""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

    def create_first_page(self, user_name, file_name):
        try:
            logger.info(f"[{self.name}] Creating first page for {user_name}")

            # First Page
            heading = self.doc.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = heading.add_run('System Requirement Specification')
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(24)

            self.doc.add_paragraph()
            self.doc.add_paragraph()

            name_paragraph = self.doc.add_paragraph()
            name_run = name_paragraph.add_run(f'Name: {user_name}')
            name_run.font.name = 'Times New Roman'
            name_run.font.size = Pt(12)

            # Add Table of Contents after the first page
            self.add_table_of_contents()

            self.save_document(file_name)
        except Exception as e:
            logger.error(f"[{self.name}] Error in create_first_page: {e}")
            raise

    def parse_markdown_text(self, text):
        """Parse markdown-style formatting from text"""
        if not isinstance(text, str):
            logger.error(f"[{self.name}] parse_markdown_text expected a string, got {type(text).__name__}")
            return []

        formatted_paragraphs = []
        paragraphs = text.split('\n\n')

        for para in paragraphs:
            if para.strip():
                heading_match = self.heading_pattern.match(para.strip())
                if heading_match:
                    formatted_paragraphs.append({
                        'type': 'heading',
                        'text': heading_match.group(1),
                        'level': para.count('#')
                    })
                    continue

                if self.bullet_pattern.search(para):
                    bullet_points = []
                    for line in para.split('\n'):
                        bullet_match = self.bullet_pattern.match(line.strip())
                        if bullet_match:
                            bullet_points.append(bullet_match.group(1))
                    formatted_paragraphs.append({
                        'type': 'bullets',
                        'points': bullet_points
                    })
                    continue

                formatted_paragraphs.append({
                    'type': 'paragraph',
                    'text': para.strip()
                })

        return formatted_paragraphs

    def add_formatted_text(self, paragraph, text):
        """Add text to paragraph with bold formatting"""
        if not isinstance(text, str):
            logger.error(f"[{self.name}] add_formatted_text expected string but got {type(text).__name__}")
            return

        remaining_text = text
        while remaining_text:
            bold_match = self.bold_pattern.search(remaining_text)
            if bold_match:
                before_bold = remaining_text[:bold_match.start()]
                if before_bold:
                    paragraph.add_run(before_bold)

                bold_text = bold_match.group(1) or bold_match.group(2)
                run = paragraph.add_run(bold_text)
                run.bold = True
                run.font.name = 'Times New Roman'

                remaining_text = remaining_text[bold_match.end():]
            else:
                paragraph.add_run(remaining_text)
                break

    def add_page(self, content, file_name):
        try:
            logger.info(f"[{self.name}] Adding page.")

            # Add page break
            self.doc.add_page_break()

            # Define section order and numbers
            section_order = [
                ('introduction', '1'),
                ('overall_description', '2'),
                ('system_features', '3'),
                ('external_interfaces', '4'),
                ('non_functional_requirements', '5'),
                ('use_cases', '6'),
            ]

            for key, section_num in section_order:
                if key not in content:
                    continue
                section_content = content[key]

                # Defensive check: section_content must be a string
                if not isinstance(section_content, str):
                    # Try to extract string if it's a dict with a single string value (common agent bug)
                    if isinstance(section_content, dict):
                        # Look for the first string value in the dict
                        str_val = next((v for v in section_content.values() if isinstance(v, str)), None)
                        if str_val:
                            logger.warning(f"[{self.name}] Section '{key}' was a dict, using first string value.")
                            section_content = str_val
                        else:
                            logger.error(f"[{self.name}] Expected string for section '{key}', got {type(section_content).__name__} instead.")
                            continue
                    else:
                        logger.error(f"[{self.name}] Expected string for section '{key}', got {type(section_content).__name__} instead.")
                        continue

                logger.info(f"[{self.name}] Writing section '{key}' with content: {repr(section_content)[:200]}")
                

                formatted_content = self.parse_markdown_text(section_content)

                # Add numbered heading
                self.doc.add_heading(f"{section_num}. {key.replace('_', ' ').title()}", level=1)

                # Only add paragraphs and bullets, skip subheadings from LLM output
                for item in formatted_content:
                    if item['type'] == 'paragraph':
                        para = self.doc.add_paragraph()
                        self.add_formatted_text(para, item['text'])
                    elif item['type'] == 'bullets':
                        for point in item['points']:
                            bullet_para = self.doc.add_paragraph(style='List Bullet')
                            self.add_formatted_text(bullet_para, point)

            self.save_document(file_name)
        except Exception as e:
            import traceback
            logger.error(f"[{self.name}] Error in add_page: {e}\n{traceback.format_exc()}")
            raise
